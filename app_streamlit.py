"""
Interface web RAG - Streamlit com historico e exportacao.

Execute com: python -m streamlit run app_streamlit.py
"""
from __future__ import annotations

import io
import sys
import html as _html
import uuid
import re
import csv
from pathlib import Path
from urllib.parse import quote

try:
    import streamlit as st
except ImportError:
    print("Erro: streamlit nao instalado.")
    print("Execute: pip install streamlit python-docx rank_bm25 pymupdf")
    sys.exit(1)

# Carrega .env do diretorio pai
import os
from dotenv import load_dotenv
load_dotenv(Path(__file__).resolve().parents[1] / ".env")

from rag_core import init_rag, answer_question, suggest_follow_up_questions

# Nomes exibidos no chat (estilo ChatPetrobras): usuário logado e assistente
CHAT_NOME_USUARIO = os.environ.get("CHAT_NOME_USUARIO") or "Nome da pessoa logada"
CHAT_NOME_ASSISTENTE = "JackChat"

# Avatar transparente 1x1 para nao exibir letra no circulo (só o nome em texto)
import base64 as _b64
_AVATAR_TRANSPARENTE = _b64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNk+M9QDwADhgGAWjR9awAAAABJRU5ErkJggg=="
)


@st.cache_resource(show_spinner=False)
def _load_rag_cached(_base_dir_str: str, _version: int, _reindex: bool = False):
    """Carrega RAG (cache por version/reindex)."""
    import os
    if _reindex:
        os.environ["REINDEX"] = "true"
    try:
        return init_rag(Path(_base_dir_str))
    finally:
        os.environ.pop("REINDEX", None)


@st.cache_data(ttl=3600)
def _img_to_b64(path: str) -> str:
    """Converte imagem para base64 (cache 1h)."""
    import base64
    return base64.b64encode(Path(path).read_bytes()).decode("utf-8")


@st.cache_resource(show_spinner=False)
def _memory_store() -> dict[str, dict]:
    # Armazenamento em memória do processo (não grava em arquivo).
    # Mantém histórico mesmo se a página recarregar, enquanto o Streamlit estiver rodando.
    return {}


def _export_docx(pergunta: str, resposta: str, fontes: list) -> bytes:
    """Gera documento Word formatado com pergunta, resposta e fontes."""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from datetime import datetime

    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    titulo = doc.add_heading("Consulta de dados contratuais", 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_data = doc.add_paragraph()
    p_data.add_run(datetime.now().strftime("%d/%m/%Y %H:%M")).italic = True
    p_data.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Pergunta:").bold = True
    p.add_run("\n" + pergunta)
    p.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph()
    p.add_run("Resposta:").bold = True
    p.add_run("\n" + resposta)
    p.paragraph_format.space_after = Pt(12)

    if fontes:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Fontes consultadas:").bold = True
        for i, f in enumerate(fontes, 1):
            doc.add_paragraph(f"{i}. {f}", style="List Number")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _export_txt(pergunta: str, resposta: str, fontes: list) -> str:
    """Gera texto simples para exportacao."""
    linhas = [
        "=== Consulta de dados contratuais ===",
        "",
        "Pergunta: " + pergunta,
        "",
        "Resposta:",
        resposta,
        "",
    ]
    if fontes:
        linhas.append("Fontes consultadas:")
        for i, f in enumerate(fontes, 1):
            linhas.append(f"  {i}. {f}")
    return "\n".join(linhas)


def _export_xlsx(pergunta: str, resposta: str, fontes: list) -> bytes:
    """Gera planilha Excel formatada com pergunta, resposta e fontes."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
    from openpyxl.utils import get_column_letter
    from datetime import datetime

    wb = Workbook()
    ws = wb.active
    ws.title = "Relatorio"

    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin"),
    )
    header_fill = PatternFill(start_color="2D5A3D", end_color="2D5A3D", fill_type="solid")

    ws["A1"] = "Consulta de dados contratuais - DGRTA"
    ws["A1"].font = Font(bold=True, size=14, color="FFFFFF")
    ws["A1"].fill = header_fill
    ws["A1"].alignment = Alignment(horizontal="center", wrap_text=True)
    ws.merge_cells("A1:C1")

    ws["A2"] = datetime.now().strftime("%d/%m/%Y %H:%M")
    ws["A2"].font = Font(italic=True, size=9)
    ws.merge_cells("A2:C2")

    row = 4
    ws[f"A{row}"] = "Pergunta:"
    ws[f"A{row}"].font = Font(bold=True)
    ws[f"A{row}"].fill = PatternFill(start_color="E8F0E8", end_color="E8F0E8", fill_type="solid")
    row += 1
    ws[f"A{row}"] = pergunta
    ws[f"A{row}"].alignment = Alignment(wrap_text=True, vertical="top")
    ws.row_dimensions[row].height = max(30, min(200, len(pergunta) // 3))
    row += 2

    ws[f"A{row}"] = "Resposta:"
    ws[f"A{row}"].font = Font(bold=True)
    ws[f"A{row}"].fill = PatternFill(start_color="E8F0E8", end_color="E8F0E8", fill_type="solid")
    row += 1
    ws[f"A{row}"] = resposta
    ws[f"A{row}"].alignment = Alignment(wrap_text=True, vertical="top")
    ws.row_dimensions[row].height = max(30, min(400, len(resposta) // 2))
    row += 2

    if fontes:
        ws[f"A{row}"] = "Fontes consultadas:"
        ws[f"A{row}"].font = Font(bold=True)
        ws[f"A{row}"].fill = PatternFill(start_color="E8F0E8", end_color="E8F0E8", fill_type="solid")
        row += 1
        for i, f in enumerate(fontes, 1):
            ws[f"A{row}"] = f"{i}. {f}"
            ws[f"A{row}"].alignment = Alignment(wrap_text=True)
            row += 1

    ws.column_dimensions["A"].width = 100
    ws.column_dimensions["B"].width = 15
    ws.column_dimensions["C"].width = 15

    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _try_parse_number_br(s: str) -> float | None:
    """Converte '1.234,56' / 'R$ 1.234,56' em float. Retorna None se falhar."""
    if s is None:
        return None
    txt = str(s).strip()
    if not txt:
        return None
    txt = txt.replace("R$", "").replace("r$", "").strip()
    txt = re.sub(r"\s+", "", txt)
    txt = re.sub(r"[^0-9,.-]", "", txt)
    if not txt:
        return None
    if "," in txt and txt.count(",") == 1:
        txt = txt.replace(".", "").replace(",", ".")
    try:
        return float(txt)
    except Exception:
        return None


def _extract_csv_table(text: str) -> list[list[str]] | None:
    """Extrai uma tabela tipo CSV da resposta do modelo (codeblock ou linhas com vírgula)."""
    if not text:
        return None

    blocks = re.findall(r"```(?:csv|text)?\n([\s\S]*?)```", text, flags=re.IGNORECASE)
    candidates: list[str] = []
    for b in blocks:
        b = b.strip()
        if b.count("\n") >= 1 and b.count(",") >= 3:
            candidates.append(b)

    if not candidates:
        lines = [ln.strip() for ln in text.splitlines()]
        best: list[str] = []
        cur: list[str] = []
        for ln in lines:
            looks_csv = ln.count(",") >= 3 and not ln.lower().startswith(("fonte", "fontes"))
            if looks_csv:
                cur.append(ln)
            else:
                if len(cur) > len(best):
                    best = cur
                cur = []
        if len(cur) > len(best):
            best = cur
        if len(best) >= 2:
            candidates.append("\n".join(best))

    if not candidates:
        return None

    csv_text = max(candidates, key=lambda s: s.count("\n"))
    try:
        reader = csv.reader(io.StringIO(csv_text), delimiter=",")
        rows = [[c.strip() for c in r] for r in reader if any(c.strip() for c in r)]
        if len(rows) >= 2 and len(rows[0]) >= 3:
            return rows
    except Exception:
        return None
    return None


def _export_xlsx_table(pergunta: str, rows: list[list[str]], fontes: list) -> bytes:
    """Gera Excel formatado a partir de uma tabela (rows[0]=cabecalho)."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
    from openpyxl.utils import get_column_letter
    from datetime import datetime

    wb = Workbook()
    ws = wb.active
    ws.title = "Tabela"

    thin = Side(style="thin")
    thin_border = Border(left=thin, right=thin, top=thin, bottom=thin)
    header_fill = PatternFill(start_color="2D5A3D", end_color="2D5A3D", fill_type="solid")

    max_cols = max(3, len(rows[0]))
    ws["A1"] = "Consulta de dados contratuais - Tabela"
    ws["A1"].font = Font(bold=True, size=14, color="FFFFFF")
    ws["A1"].fill = header_fill
    ws["A1"].alignment = Alignment(horizontal="center")
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=max_cols)

    ws["A2"] = datetime.now().strftime("%d/%m/%Y %H:%M")
    ws["A2"].font = Font(italic=True, size=9)
    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=max_cols)

    ws["A4"] = "Pergunta:"
    ws["A4"].font = Font(bold=True)
    ws["A5"] = pergunta
    ws["A5"].alignment = Alignment(wrap_text=True, vertical="top")
    ws.merge_cells(start_row=5, start_column=1, end_row=5, end_column=max_cols)

    start_row = 7
    header = rows[0]
    for c_idx, col in enumerate(header, 1):
        cell = ws.cell(row=start_row, column=c_idx, value=col)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = thin_border

    norm = [re.sub(r"\s+", " ", h.strip().lower()) for h in header]
    num_cols: set[int] = set()
    for i, h in enumerate(norm, 1):
        if any(k in h for k in ("quantidade", "preço", "preco", "valor", "unitário", "unitario", "total", "r$")):
            num_cols.add(i)

    for r_idx, row in enumerate(rows[1:], start_row + 1):
        for c_idx in range(1, len(header) + 1):
            val = row[c_idx - 1] if c_idx - 1 < len(row) else ""
            cell = ws.cell(row=r_idx, column=c_idx, value=val)
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = thin_border
            if c_idx in num_cols:
                n = _try_parse_number_br(val)
                if n is not None:
                    cell.value = n
                    cell.number_format = '#,##0.00'
                    cell.alignment = Alignment(vertical="top", horizontal="right")

    last_row = start_row + len(rows) - 1
    last_col = len(header)
    ws.freeze_panes = ws["A8"]
    ws.auto_filter.ref = f"A{start_row}:{get_column_letter(last_col)}{last_row}"

    for c_idx in range(1, last_col + 1):
        col_letter = get_column_letter(c_idx)
        max_len = 10
        for r in range(start_row, min(last_row, start_row + 200) + 1):
            v = ws.cell(row=r, column=c_idx).value
            if v is None:
                continue
            max_len = max(max_len, len(str(v)))
        ws.column_dimensions[col_letter].width = min(60, max(12, max_len + 2))

    if fontes:
        ws2 = wb.create_sheet("Fontes")
        ws2["A1"] = "Fontes consultadas"
        ws2["A1"].font = Font(bold=True)
        for i, f in enumerate(fontes, 1):
            ws2[f"A{i+1}"] = f

    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _escape_html(s: str) -> str:
    """Escapa caracteres especiais para uso em HTML/ReportLab."""
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _export_pdf(pergunta: str, resposta: str, fontes: list) -> bytes:
    """Gera PDF com pergunta, resposta e fontes."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=2*cm, leftMargin=2*cm)
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph("Consulta de dados contratuais", styles["Title"]))
    story.append(Spacer(1, 0.5*cm))
    story.append(Paragraph("<b>Pergunta:</b>", styles["Normal"]))
    story.append(Paragraph(_escape_html(pergunta).replace("\n", "<br/>"), styles["Normal"]))
    story.append(Spacer(1, 0.5*cm))
    story.append(Paragraph("<b>Resposta:</b>", styles["Normal"]))
    story.append(Paragraph(_escape_html(resposta).replace("\n", "<br/>"), styles["Normal"]))
    if fontes:
        story.append(Spacer(1, 0.5*cm))
        story.append(Paragraph("<b>Fontes consultadas:</b>", styles["Normal"]))
        for i, f in enumerate(fontes, 1):
            story.append(Paragraph(f"{i}. {_escape_html(f)}", styles["Normal"]))

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


def _render_sidebar_history(placeholder, messages: list[dict]) -> None:
    with placeholder.container():
        st.subheader("📋 Histórico de perguntas")
        perguntas = [m["content"] for m in messages if m.get("role") == "user"]
        if perguntas:
            # Ultimas 5 perguntas (limite do historico)
            ultimas = perguntas[-5:]
            for i, p in enumerate(ultimas, 1):
                texto = (p[:60] + "...") if len(p) > 60 else p
                if st.button(f"{i}. {texto}", key=f"hist_{i}_{hash(p)}", use_container_width=True):
                    st.session_state._hist_click = p
                    st.rerun()
        else:
            st.caption("Nenhuma pergunta ainda.")


def main() -> None:
    st.set_page_config(
        page_title="Consulta de dados contratuais",
        page_icon="📄",
        layout="centered",
        initial_sidebar_state="expanded",
    )

    # Layout: sidebar branca, faixas verdes
    st.markdown(
        """
        <style>
        /* Esconde header padrao do Streamlit - NAO esconder collapsedControl (seta para expandir sidebar) */
        #MainMenu, footer { visibility: hidden !important; }
        header[data-testid="stHeader"] { visibility: hidden !important; height: 0 !important; min-height: 0 !important; overflow: visible !important; }
        header { height: 0 !important; min-height: 0 !important; overflow: visible !important; }
        /* Garante que o botao de expandir sidebar fique sempre visivel quando recolhida */
        [data-testid="collapsedControl"],
        [data-testid="stSidebar"] [data-testid="collapsedControl"] {
            visibility: visible !important; display: flex !important; opacity: 1 !important; z-index: 9999 !important;
        }
        
        .stApp { background-color: #f5f5f5 !important; }
        
        /* Sidebar verde mais claro (mesmo tom da barra superior) - fixa na lateral ao rolar */
        [data-testid="stSidebar"],
        [data-testid="stSidebar"] > div:first-child,
        section[data-testid="stSidebar"] {
            background: #5d8a6d !important;
            border-right: 1px solid #4a7a55 !important;
            width: 21rem !important;
            max-width: 21rem !important;
            position: sticky !important;
            top: 0 !important;
            align-self: start !important;
            overflow-y: auto !important;
        }
        [data-testid="stSidebar"] .stMarkdown,
        [data-testid="stSidebar"] h1, [data-testid="stSidebar"] h2, [data-testid="stSidebar"] h3,
        [data-testid="stSidebar"] .stCaption,
        [data-testid="stSidebar"] label { color: #ffffff !important; }
        [data-testid="stSidebar"] .stButton button {
            color: #2d5a3d !important; background: #ffffff !important;
        }
        /* Substitui "Drag and drop files here" por "Inclua o PDF aqui" */
        [data-testid="stSidebar"] [data-testid="stFileUploader"] small,
        [data-testid="stSidebar"] [data-testid="stFileUploader"] p,
        [data-testid="stSidebar"] [data-testid="stFileUploader"] [data-testid="stFileUploaderDropzoneText"],
        [data-testid="stSidebar"] .stFileUploader small,
        [data-testid="stSidebar"] .stFileUploader p,
        [data-testid="stSidebar"] div[data-testid="stFileUploader"] > div > div > small,
        [data-testid="stSidebar"] div[data-testid="stFileUploader"] > div > div > p {
            display: none !important;
            visibility: hidden !important;
            height: 0 !important;
            overflow: hidden !important;
            margin: 0 !important;
            padding: 0 !important;
            font-size: 0 !important;
            line-height: 0 !important;
        }
        [data-testid="stSidebar"] [data-testid="stFileUploader"] > div > div:first-child::before,
        [data-testid="stSidebar"] [data-testid="stFileUploader"] [data-testid="stFileUploaderDropzone"]::before {
            content: "Inclua o PDF aqui";
            display: block !important;
            text-align: center;
            color: #ffffff !important;
            padding: 0.5rem 0;
            font-size: 0.9rem;
        }
        
        /* Area principal - espaco extra para area fixa de sugestoes + input */
        main .block-container { 
            padding: 1rem !important;
            padding-bottom: 320px !important; 
            background-color: #ffffff !important;
            border-radius: 0 !important;
            box-shadow: 0 1px 3px rgba(0,0,0,0.08);
            margin-left: 0 !important;
        }
        
        /* Barra de navegacao verde - fixa abaixo do banner */
        .nav-bar-dgrta {
            position: sticky !important;
            top: 0 !important;
            background: #2d5a3d !important;
            padding: 0.5rem 1rem !important;
            margin: -1rem -1rem 1rem -1rem !important;
            display: flex !important;
            gap: 1.5rem !important;
            align-items: center !important;
            z-index: 997 !important;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }
        .nav-bar-dgrta span { color: white !important; font-size: 0.95rem !important; cursor: pointer !important; }
        
        :root {
            --dgrta-sidebar-w: 21rem;
            --dgrta-chatbar-h: 130px; /* altura real aproximada do chat_input (evita sobreposicao) */
            --dgrta-gap: 0px;
            --dgrta-dock-bg: #e8e8e8;
        }

        /* Rodape fixo: sugestoes sempre acima do input (uma linha) */
        #dgrta-sugestoes-dock {
            position: fixed !important;
            bottom: calc(var(--dgrta-chatbar-h) + var(--dgrta-gap)) !important; /* imediatamente acima do st.chat_input */
            left: var(--dgrta-sidebar-w) !important;   /* NAO cobrir a sidebar */
            right: 0 !important;
            z-index: 998 !important;
            background: var(--dgrta-dock-bg) !important;
            padding: 0.4rem 0.75rem !important;
            box-shadow: 0 -2px 10px rgba(0,0,0,0.08);
            display: flex !important;
            align-items: center !important;
            gap: 0.75rem !important;
        }
        .dgrta-sug-label {
            font-weight: 600;
            color: #2d5a3d;
            white-space: nowrap;
            flex: 0 0 auto;
            pointer-events: auto !important;
        }
        #dgrta-sug-row {
            display: flex;
            align-items: center;
            gap: 0.5rem;
            overflow-x: auto;
            overflow-y: hidden;
            white-space: nowrap;
            flex: 1 1 auto;
            padding-bottom: 0.15rem;
            pointer-events: auto !important;
        }
        .dgrta-sug-btn {
            display: inline-block;
            background: #ffffff;
            border: 1px solid #dcdcdc;
            border-radius: 999px;
            padding: 0.45rem 0.7rem;
            color: #2d5a3d;
            text-decoration: none;
            text-align: center;
            font-size: 0.9rem;
            line-height: 1.2;
            white-space: nowrap;
            flex: 0 0 auto;
            cursor: pointer;
        }
        .dgrta-sug-btn:hover {
            border-color: #2d5a3d;
        }
        /* Input fixo no rodape (garante consistencia entre navegadores) */
        [data-testid="stChatInputContainer"] {
            position: fixed !important;
            bottom: 0 !important;
            left: var(--dgrta-sidebar-w) !important; /* NAO cobrir a sidebar */
            right: 0 !important;
            z-index: 999 !important;
            background: var(--dgrta-dock-bg) !important;
            padding: 0.5rem 1rem 0.75rem !important;
            box-shadow: 0 -2px 10px rgba(0,0,0,0.1);
            min-height: var(--dgrta-chatbar-h) !important;
            box-sizing: border-box !important;
        }
        [data-testid="stChatInputContainer"] * { visibility: visible !important; }

        /* Responsivo: em telas menores, ocupar a largura toda */
        @media (max-width: 900px) {
            :root { --dgrta-sidebar-w: 0rem; }
        }
        
        /* Esconde só o avatar (não mexe na largura das colunas para não quebrar o texto) */
        [data-testid="stChatMessageAvatar"],
        [data-testid="stChatMessage"] [data-testid="stChatMessageAvatar"] {
            display: none !important;
        }
        /* Nome do chat em destaque */
        .dgrta-chat-nome {
            font-weight: bold;
            font-size: 0.95rem;
            margin-bottom: 0.25rem;
            color: inherit !important;
        }
        .dgrta-chat-nome.user { text-align: right; }
        .dgrta-chat-nome.assistant { text-align: left; }

        /* Fundo para separar perguntas das respostas */
        [data-testid="stChatMessage"]:has(.dgrta-chat-nome.user) {
            background: #f0f4f0 !important;
            border-radius: 10px !important;
            padding: 0.75rem 1rem !important;
            margin-bottom: 0.5rem !important;
            border-left: 4px solid #5d8a6d !important;
        }
        [data-testid="stChatMessage"]:has(.dgrta-chat-nome.assistant) {
            background: #ffffff !important;
            border-radius: 10px !important;
            padding: 0.75rem 1rem !important;
            margin-bottom: 0.5rem !important;
            box-shadow: 0 1px 3px rgba(0,0,0,0.08) !important;
            border-left: 4px solid #2d5a3d !important;
        }

        /* Banner */
        #banner-dgrta {
            width: calc(100vw - 21rem) !important;
            margin-left: calc(-1 * (100vw - 21rem - 100%) / 2) !important;
            max-width: none !important;
            box-sizing: border-box !important;
        }
        #banner-dgrta img {
            width: 100% !important;
            height: auto !important;
            object-fit: contain !important;
            object-position: center !important;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    base_dir = Path(__file__).resolve().parents[1]
    banner_path = base_dir / "assets" / "banner_dgrta.png"
    if banner_path.exists():
        b64 = _img_to_b64(str(banner_path))
        st.markdown(
            f"""
            <div id="banner-dgrta" style="
                margin-bottom: 1rem;
                border-radius: 8px;
            ">
                <img src="data:image/png;base64,{b64}" 
                     alt="DGRTA" 
                     style="
                         width: 100%;
                         height: auto;
                         display: block;
                         object-fit: contain;
                         border-radius: 8px;
                     " />
                <div style="
                    text-align: center;
                    padding: 0.75rem 1rem;
                    color: #1a472a;
                    font-size: 1.75rem;
                    font-weight: bold;
                    margin: 0;
                    line-height: 1.4;
                ">Consulta de dados contratuais</div>
            </div>
            """,
            unsafe_allow_html=True,
        )
    else:
        st.title("Consulta de dados contratuais")
    
    # Barra de navegacao verde clara (estilo SharePoint)
    st.markdown(
        """
        <div class="nav-bar-dgrta">
            <span><b>DGRTA</b></span>
        </div>
        """,
        unsafe_allow_html=True,
    )
    st.caption("Chat para perguntas e respostas baseadas em arquivos/documentos dos contratos da DGRTA.")

    # Sugestoes padrao (quando nao ha pergunta anterior)
    SUGESTOES_PADRAO = [
        "Quais são as partes do contrato?",
        "Qual o valor total do contrato?",
        "Quais são os prazos de vigência?",
        "Quais são as obrigações do contratado?",
        "Quais ICJs estão previstos no contrato?",
        "Quais são as cláusulas de penalidade?",
        "Qual o objeto do contrato?",
        "Quais documentos compõem o contrato?",
    ]
    if "sugestao_clicada" not in st.session_state:
        st.session_state.sugestao_clicada = None
    if "follow_up_suggestions" not in st.session_state:
        st.session_state.follow_up_suggestions = None

    pdf_dir = base_dir / "Base"

    # Inicializa historico e estado
    if "messages" not in st.session_state:
        st.session_state.messages = []
    if "rag_version" not in st.session_state:
        st.session_state.rag_version = 0
    if "sid" not in st.session_state:
        st.session_state.sid = None

    # Session-id no URL para não perder histórico ao clicar nas sugestões
    sid: str | None = None
    try:
        sid_raw = st.query_params.get("sid")
        sid = sid_raw[0] if isinstance(sid_raw, list) else sid_raw
    except Exception:
        try:
            qp = st.experimental_get_query_params()
            sid_list = qp.get("sid")
            sid = sid_list[0] if sid_list else None
        except Exception:
            sid = None

    if not sid:
        sid = st.session_state.sid or str(uuid.uuid4())
        st.session_state.sid = sid
        try:
            st.query_params["sid"] = sid
        except Exception:
            try:
                st.experimental_set_query_params(sid=sid)
            except Exception:
                pass
    else:
        st.session_state.sid = sid

    store = _memory_store()
    if sid not in store:
        store[sid] = {"messages": [], "follow_up_suggestions": None, "_last_sug_for": None}

    # Se a sessão recarregou, restaura histórico da memória
    if not st.session_state.messages and store[sid].get("messages"):
        st.session_state.messages = list(store[sid]["messages"])
        st.session_state.follow_up_suggestions = store[sid].get("follow_up_suggestions")
        if store[sid].get("_last_sug_for") is not None:
            st.session_state["_last_sug_for"] = store[sid].get("_last_sug_for")

    # Carrega RAG no inicio (processamento antecipado, respostas rapidas)
    _reindex = st.session_state.pop("reindex_needed", False)
    try:
        with st.spinner("Carregando documentos..."):
            rag = _load_rag_cached(str(base_dir), st.session_state["rag_version"], _reindex)
    except Exception as e:
        rag = None
        st.error(f"Erro ao carregar RAG: {e}")
        st.info("Verifique o arquivo .env e se os PDFs estão na pasta Base.")

    # Sidebar
    with st.sidebar:
        st.header("Configurações")

        # --- Incluir contratos ---
        st.subheader("📤 Incluir contratos")
        uploaded = st.file_uploader("Enviar novos PDFs", type=["pdf"], accept_multiple_files=True, key="upload_contratos")
        if uploaded:
            pdf_dir.mkdir(parents=True, exist_ok=True)
            for f in uploaded:
                path = pdf_dir / f.name
                path.write_bytes(f.getvalue())
            st.session_state.rag_version += 1
            st.session_state.reindex_needed = True
            st.success(f"{len(uploaded)} arquivo(s) salvo(s). Reindexando...")
            st.rerun()
        st.caption("Adicione PDFs dos contratos para consulta.")

        st.divider()

        sidebar_hist_ph = st.empty()
        _render_sidebar_history(sidebar_hist_ph, st.session_state.messages)

        st.divider()

        # --- Limpar chat ---
        st.subheader("🗑️ Chat")
        if st.button("🔄 Limpar histórico", use_container_width=True, key="btn_limpar"):
            st.session_state.messages = []
            st.session_state.follow_up_suggestions = None
            try:
                store[sid]["messages"] = []
                store[sid]["follow_up_suggestions"] = None
                store[sid]["_last_sug_for"] = None
            except Exception:
                pass
            st.rerun()

        st.divider()

        # Filtrar por documento
        st.subheader("📁 Filtrar por documento")
        pdf_files = sorted(pdf_dir.glob("**/*.pdf")) if pdf_dir.exists() else []
        opcoes = ["Todos os documentos"] + [p.name for p in pdf_files]
        filter_doc = st.selectbox("Buscar apenas em:", opcoes, key="filter_doc")
        doc_filter = None if filter_doc == "Todos os documentos" else filter_doc

    # Exibe historico de mensagens (nomes estilo ChatPetrobras: usuário logado e JackChat)
    for i, msg in enumerate(st.session_state.messages):
        nome_chat = CHAT_NOME_USUARIO if msg["role"] == "user" else CHAT_NOME_ASSISTENTE
        with st.chat_message(nome_chat, avatar=_AVATAR_TRANSPARENTE):
            st.markdown(f'<p class="dgrta-chat-nome {msg["role"]}">{_html.escape(nome_chat)}</p>', unsafe_allow_html=True)
            st.markdown(msg["content"])
            if msg.get("sources"):
                with st.expander("📎 Fontes"):
                    for s in msg["sources"]:
                        st.write(f"• {s}")

    # Sugestoes: padrao se sem pergunta; contextual com base na ultima pergunta/resposta
    msgs = st.session_state.messages
    tem_resposta = len(msgs) >= 2 and msgs[-1]["role"] == "assistant"
    if tem_resposta:
        ultima_pergunta = msgs[-2]["content"]
        ultima_resposta = msgs[-1]["content"]
        sugestoes_atuais = st.session_state.follow_up_suggestions
        if sugestoes_atuais is None or st.session_state.get("_last_sug_for") != (ultima_pergunta[:50], len(msgs)):
            sugestoes_atuais = SUGESTOES_PADRAO  # fallback ate gerar (ou se LLM falhou)
    else:
        sugestoes_atuais = SUGESTOES_PADRAO

    # Sugestoes fixas no rodape (HTML). Clique -> ?sug=... -> vira prompt.
    sug_from_url: str | None = None
    try:
        if "sug" in st.query_params:
            raw = st.query_params.get("sug")
            sug_from_url = raw[0] if isinstance(raw, list) else raw
            # Remove apenas "sug", preserva "sid"
            try:
                st.query_params.clear()
                st.query_params["sid"] = sid
            except Exception:
                pass
    except Exception:
        try:
            qp = st.experimental_get_query_params()
            if "sug" in qp and qp["sug"]:
                sug_from_url = qp["sug"][0]
                st.experimental_set_query_params(sid=sid)
        except Exception:
            sug_from_url = None

    # Link HTML (sem JS): dispara ?sug=... na MESMA aba.
    # (JS em onclick pode ser bloqueado pelo Streamlit/CSP em alguns navegadores)
    sug_links = "\n".join(
        f'<a class="dgrta-sug-btn" target="_self" href="?sid={quote(sid)}&sug={quote(s)}">{_html.escape(s)}</a>'
        for s in sugestoes_atuais[:8]
    )
    st.markdown(
        f"""
        <div id="dgrta-sugestoes-dock">
          <span class="dgrta-sug-label">💡 Sugestões:</span>
          <div id="dgrta-sug-row">
            {sug_links}
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    chat_value = st.chat_input("Digite sua pergunta sobre os contratos...")
    hist_click = st.session_state.pop("_hist_click", None)
    prompt = hist_click or sug_from_url or chat_value

    if prompt:
        st.session_state.messages.append({"role": "user", "content": prompt})
        try:
            store[sid]["messages"] = list(st.session_state.messages)
        except Exception:
            pass

        with st.chat_message(CHAT_NOME_USUARIO, avatar=_AVATAR_TRANSPARENTE):
            st.markdown(f'<p class="dgrta-chat-nome user">{_html.escape(CHAT_NOME_USUARIO)}</p>', unsafe_allow_html=True)
            st.markdown(prompt)

        if not rag:
            with st.chat_message(CHAT_NOME_ASSISTENTE, avatar=_AVATAR_TRANSPARENTE):
                st.markdown(f'<p class="dgrta-chat-nome assistant">{_html.escape(CHAT_NOME_ASSISTENTE)}</p>', unsafe_allow_html=True)
                st.error("Não foi possível carregar o sistema. Verifique o arquivo .env e os PDFs na pasta Base.")
            st.session_state.messages.append({
                "role": "assistant",
                "content": "Não foi possível carregar o sistema. Verifique o arquivo .env e os PDFs na pasta Base.",
                "sources": [],
            })
        else:
            with st.chat_message(CHAT_NOME_ASSISTENTE, avatar=_AVATAR_TRANSPARENTE):
                st.markdown(f'<p class="dgrta-chat-nome assistant">{_html.escape(CHAT_NOME_ASSISTENTE)}</p>', unsafe_allow_html=True)
                with st.status("⏳ Processando sua pergunta...", expanded=True):
                    st.write("1. Buscando trechos relevantes nos documentos...")
                    msgs = st.session_state.messages
                    history = [
                        (msgs[i]["content"], msgs[i + 1]["content"])
                        for i in range(0, len(msgs) - 1, 2)
                        if i + 1 < len(msgs) and msgs[i]["role"] == "user" and msgs[i + 1]["role"] == "assistant"
                    ]
                    answer, sources = answer_question(
                        prompt,
                        rag["retriever"],
                        rag["llm"],
                        rag["pdf_dir"],
                        history=history,
                        filter_source=doc_filter,
                        vectorstore=rag.get("vectorstore"),
                    )

                st.markdown(answer)

                source_labels = [rag["get_source_label"](d) for d in sources]
                if source_labels:
                    with st.expander("📎 Fontes consultadas"):
                        for s in source_labels:
                            st.write(f"• {s}")

                # Botoes de exportacao (TXT, Word, Excel, PDF)
                col1, col2, col3, col4 = st.columns(4)
                export_key = f"export_{len(st.session_state.messages)}"

                with col1:
                    txt_data = _export_txt(prompt, answer, source_labels)
                    st.download_button(
                        "📝 TXT",
                        data=txt_data,
                        file_name="consulta_rag.txt",
                        mime="text/plain",
                        key=f"{export_key}_txt",
                    )
                with col2:
                    try:
                        docx_bytes = _export_docx(prompt, answer, source_labels)
                        st.download_button(
                            "📘 Word",
                            data=docx_bytes,
                            file_name="consulta_rag.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"{export_key}_docx",
                        )
                    except ImportError:
                        st.caption("python-docx não instalado")
                with col3:
                    try:
                        rows = _extract_csv_table(answer)
                        if rows is not None:
                            xlsx_bytes = _export_xlsx_table(prompt, rows, source_labels)
                            file_name = "consulta_tabela.xlsx"
                        else:
                            xlsx_bytes = _export_xlsx(prompt, answer, source_labels)
                            file_name = "consulta_rag.xlsx"
                        st.download_button(
                            "📊 Excel",
                            data=xlsx_bytes,
                            file_name=file_name,
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            key=f"{export_key}_xlsx",
                        )
                    except ImportError:
                        st.caption("openpyxl não instalado")
                with col4:
                    try:
                        pdf_bytes = _export_pdf(prompt, answer, source_labels)
                        st.download_button(
                            "📕 PDF",
                            data=pdf_bytes,
                            file_name="consulta_rag.pdf",
                            mime="application/pdf",
                            key=f"{export_key}_pdf",
                        )
                    except ImportError:
                        st.caption("reportlab não instalado")

            st.session_state.messages.append({
                "role": "assistant",
                "content": answer,
                "sources": source_labels,
            })
            # Limita historico a 5 trocas (10 mensagens: 5 user + 5 assistant)
            if len(st.session_state.messages) > 10:
                st.session_state.messages = st.session_state.messages[-10:]
            try:
                store[sid]["messages"] = list(st.session_state.messages)
            except Exception:
                pass

            # Atualiza histórico na sidebar no mesmo ciclo (sem rerun)
            try:
                _render_sidebar_history(sidebar_hist_ph, st.session_state.messages)
            except Exception:
                pass
            # Gera sugestoes de acompanhamento com base na pergunta e resposta
            try:
                follow_ups = suggest_follow_up_questions(prompt, answer, rag["llm"])
                if follow_ups:
                    st.session_state.follow_up_suggestions = follow_ups
                    st.session_state._last_sug_for = (prompt[:50], len(st.session_state.messages))
                    try:
                        store[sid]["follow_up_suggestions"] = follow_ups
                        store[sid]["_last_sug_for"] = st.session_state._last_sug_for
                    except Exception:
                        pass
            except Exception:
                st.session_state.follow_up_suggestions = SUGESTOES_PADRAO
                try:
                    store[sid]["follow_up_suggestions"] = SUGESTOES_PADRAO
                    store[sid]["_last_sug_for"] = None
                except Exception:
                    pass


if __name__ == "__main__":
    main()
